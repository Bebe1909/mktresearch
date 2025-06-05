#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Comprehensive Word Export for Layered Research Results
Hỗ trợ xuất báo cáo với Layer 3 và comprehensive Layer 4
Enhanced with tables and executive summary
"""

import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import re

def get_vietnamese_market_name(market: str) -> str:
    """Translate market name to Vietnamese for consistent Vietnamese reports"""
    market_translations = {
        "🇺🇸 United States": "🇺🇸 Hoa Kỳ",
        "🇨🇳 China": "🇨🇳 Trung Quốc", 
        "🇯🇵 Japan": "🇯🇵 Nhật Bản",
        "🇰🇷 South Korea": "🇰🇷 Hàn Quốc",
        "🇹🇭 Thailand": "🇹🇭 Thái Lan",
        "🇸🇬 Singapore": "🇸🇬 Singapore",
        "🇲🇾 Malaysia": "🇲🇾 Malaysia",
        "🇮🇩 Indonesia": "🇮🇩 Indonesia",
        "🇵🇭 Philippines": "🇵🇭 Philippines",
        "🇬🇧 United Kingdom": "🇬🇧 Vương quốc Anh",
        "🇩🇪 Germany": "🇩🇪 Đức",
        "🇫🇷 France": "🇫🇷 Pháp",
        "🇮🇹 Italy": "🇮🇹 Ý",
        "🇪🇸 Spain": "🇪🇸 Tây Ban Nha",
        "🇨🇦 Canada": "🇨🇦 Canada",
        "🇦🇺 Australia": "🇦🇺 Úc",
        "🇳🇿 New Zealand": "🇳🇿 New Zealand",
        "🇧🇷 Brazil": "🇧🇷 Brazil",
        "🇲🇽 Mexico": "🇲🇽 Mexico",
        "🇮🇳 India": "🇮🇳 Ấn Độ",
        "🇷🇺 Russia": "🇷🇺 Nga",
        "🇿🇦 South Africa": "🇿🇦 Nam Phi",
        "🇪🇬 Egypt": "🇪🇬 Ai Cập",
        "🇦🇪 UAE": "🇦🇪 UAE",
        "🇸🇦 Saudi Arabia": "🇸🇦 Ả Rập Saudi",
        "🇹🇷 Turkey": "🇹🇷 Thổ Nhĩ Kỳ",
        "🇳🇱 Netherlands": "🇳🇱 Hà Lan",
        "🇸🇪 Sweden": "🇸🇪 Thụy Điển",
        "🇳🇴 Norway": "🇳🇴 Na Uy",
        "🇩🇰 Denmark": "🇩🇰 Đan Mạch",
        "🇫🇮 Finland": "🇫🇮 Phần Lan",
        "🇨🇭 Switzerland": "🇨🇭 Thụy Sĩ",
        "🇦🇹 Austria": "🇦🇹 Áo",
        "🇧🇪 Belgium": "🇧🇪 Bỉ",
        "🇵🇱 Poland": "🇵🇱 Ba Lan",
        "🇨🇿 Czech Republic": "🇨🇿 Cộng hòa Séc",
        "🇭🇺 Hungary": "🇭🇺 Hungary",
        "🇬🇷 Greece": "🇬🇷 Hy Lạp",
        "🇵🇹 Portugal": "🇵🇹 Bồ Đào Nha",
        "🇮🇪 Ireland": "🇮🇪 Ireland",
        "🇮🇱 Israel": "🇮🇱 Israel",
        "🇭🇰 Hong Kong": "🇭🇰 Hồng Kông",
        "🇹🇼 Taiwan": "🇹🇼 Đài Loan",
        "🇦🇷 Argentina": "🇦🇷 Argentina",
        "🇨🇱 Chile": "🇨🇱 Chile",
        "🇨🇴 Colombia": "🇨🇴 Colombia",
        "🇵🇪 Peru": "🇵🇪 Peru",
        "🇻🇪 Venezuela": "🇻🇪 Venezuela",
        "🇪🇨 Ecuador": "🇪🇨 Ecuador",
        "🇺🇾 Uruguay": "🇺🇾 Uruguay",
        "🇧🇴 Bolivia": "🇧🇴 Bolivia",
        "🇵🇾 Paraguay": "🇵🇾 Paraguay",
        "🇳🇬 Nigeria": "🇳🇬 Nigeria",
        "🇰🇪 Kenya": "🇰🇪 Kenya",
        "🇬🇭 Ghana": "🇬🇭 Ghana",
        "🇪🇹 Ethiopia": "🇪🇹 Ethiopia",
        "🇺🇬 Uganda": "🇺🇬 Uganda",
        "🇹🇿 Tanzania": "🇹🇿 Tanzania",
        "🇿🇼 Zimbabwe": "🇿🇼 Zimbabwe",
        "🌏 Southeast Asia": "🌏 Đông Nam Á",
        "🌍 Asia-Pacific": "🌍 Châu Á - Thái Bình Dương",
        "🌎 Global Market": "🌎 Thị trường Toàn cầu"
    }
    
    return market_translations.get(market, market)

def add_custom_styles(doc):
    """Thêm custom styles cho document"""
    
    # Style cho tiêu đề chính
    try:
        title_style = doc.styles.add_style('CustomTitle', 1)  # 1 = PARAGRAPH
        title_font = title_style.font
        title_font.name = 'Times New Roman'
        title_font.size = Pt(20)
        title_font.bold = True
        title_font.color.rgb = RGBColor(0, 51, 102)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(20)
    except:
        pass  # Style đã tồn tại
    
    # Style cho layer heading
    try:
        layer_style = doc.styles.add_style('LayerStyle', 1)
        layer_font = layer_style.font
        layer_font.name = 'Times New Roman'
        layer_font.size = Pt(16)
        layer_font.bold = True
        layer_font.color.rgb = RGBColor(46, 134, 171)
        layer_style.paragraph_format.space_before = Pt(15)
        layer_style.paragraph_format.space_after = Pt(10)
    except:
        pass

def set_paragraph_font(paragraph, font_name='Times New Roman', font_size=11):
    """Set font cho paragraph"""
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)

def create_info_table(doc, data):
    """Tạo bảng thông tin tổng quan"""
    print("📊 Tạo bảng thông tin tổng quan...")
    
    # Tính toán statistics
    stats = calculate_statistics(data)
    
    # Get metadata from research_metadata (correct location)
    metadata = data.get('research_metadata', {})
    
    # Tạo table 2 cột
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Thêm rows - get data from metadata
    info_data = [
        ('🎯 Ngành nghiên cứu', metadata.get('industry', 'N/A')),
        ('🌍 Thị trường', metadata.get('market', 'N/A')),
        ('🤖 AI Engine', f"{metadata.get('api_provider', 'N/A')} - {metadata.get('model_used', 'N/A')}"),
        ('📅 Ngày tạo', datetime.now().strftime('%d/%m/%Y %H:%M')),
        ('❓ Tổng số questions', str(stats['total_questions']))
    ]
    
    for label, value in info_data:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        
        # Format cells
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                set_paragraph_font(paragraph, font_size=10)
                if cell == row.cells[0]:  # Label cell
                    paragraph.runs[0].bold = True
    
    return table

def create_summary_table(doc, data):
    """Tạo bảng tóm tắt kết quả nghiên cứu"""
    print("📊 Tạo bảng tóm tắt kết quả...")
    
    # Tạo table với headers
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    headers = ['Layer/Category', 'Questions', 'Layer 3 Analysis', 'Layer 4 Enhanced']
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        header_row.cells[i].text = header
        # Bold headers
        for paragraph in header_row.cells[i].paragraphs:
            set_paragraph_font(paragraph, font_size=10)
            paragraph.runs[0].bold = True
    
    # Add data rows
    for layer in data.get('research_results', []):
        layer_name = layer.get('layer_name', '')
        
        for category in layer.get('categories', []):
            category_name = category.get('category_name', '')
            questions = category.get('questions', [])
            
            layer3_count = len(questions)
            layer4_count = sum(1 for q in questions if q.get('layer4_comprehensive_report'))
            
            row = table.add_row()
            row.cells[0].text = f"{layer_name} / {category_name}"
            row.cells[1].text = str(layer3_count)
            row.cells[2].text = "✅" if layer3_count > 0 else "❌"
            row.cells[3].text = f"✅ ({layer4_count})" if layer4_count > 0 else "❌"
            
            # Format cells
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    set_paragraph_font(paragraph, font_size=9)
    
    return table

def extract_key_insights(data):
    """Trích xuất key insights để tạo executive summary"""
    insights = []
    
    for layer in data.get('research_results', []):
        layer_name = layer.get('layer_name', '')
        
        for category in layer.get('categories', []):
            category_name = category.get('category_name', '')
            
            for question in category.get('questions', []):
                # Lấy từ comprehensive report trước
                layer4_comprehensive = question.get('layer4_comprehensive_report', {})
                if layer4_comprehensive:
                    content = layer4_comprehensive.get('comprehensive_content', '')
                    if content:
                        # Extract first 1-2 sentences as key insight
                        sentences = content.split('.')[:2]
                        if sentences:
                            insight = '. '.join(sentences).strip()
                            if len(insight) > 50:  # Only meaningful insights
                                insights.append({
                                    'category': f"{layer_name} - {category_name}",
                                    'insight': insight[:200] + "..." if len(insight) > 200 else insight
                                })
                # Fallback to layer 3
                elif question.get('layer3_content'):
                    content = question.get('layer3_content', '')
                    sentences = content.split('.')[:1]
                    if sentences:
                        insight = sentences[0].strip()
                        if len(insight) > 50:
                            insights.append({
                                'category': f"{layer_name} - {category_name}",
                                'insight': insight[:150] + "..." if len(insight) > 150 else insight
                            })
    
    return insights[:8]  # Top 8 insights

def create_references_section(doc, data):
    """Tạo phần references cho báo cáo dựa trên nguồn thực từ research"""
    
    # Page break before references
    doc.add_page_break()
    
    # Title
    ref_title = doc.add_heading('📚 TÀI LIỆU THAM KHẢO', level=1)
    ref_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(ref_title, font_size=16)
    
    # Add spacing
    doc.add_paragraph()
    
    # Get research metadata
    metadata = data.get('research_metadata', {})
    topic = metadata.get('industry', 'Nghiên cứu thị trường')
    market = metadata.get('market', 'Việt Nam')
    vietnamese_market = get_vietnamese_market_name(market)  # Translate to Vietnamese
    model_used = metadata.get('model_used', 'gpt-3.5-turbo')
    api_provider = metadata.get('api_provider', 'OpenAI')
    current_year = datetime.now().year
    
    # Start with AI source acknowledgment
    references = [
        f"1. {api_provider} {model_used}. ({current_year}). AI-powered market research analysis for {topic} in {vietnamese_market}. Retrieved from https://openai.com"
    ]
    
    # Get tracked references from research data
    tracked_references = data.get('tracked_references', [])
    
    if tracked_references:
        # Use real tracked references
        ref_counter = 2
        for source, frequency in tracked_references[:10]:  # Top 10 most cited
            # Clean up source name - remove extra text that might be added by tracking
            clean_source = source.strip()
            
            # Create professional reference format
            references.append(f"{ref_counter}. {clean_source}. ({current_year}). Market research data and analysis. Industry intelligence source.")
            ref_counter += 1
    else:
        # Fallback to generic sources if no tracking data
        fallback_sources = [
            f"2. General Statistics Office (GSO). ({current_year}). Economic and social statistics. Retrieved from https://gso.gov.vn",
            f"3. World Bank. ({current_year}). World Development Indicators. Retrieved from https://data.worldbank.org",
            f"4. Vietnam Chamber of Commerce and Industry (VCCI). ({current_year}). Business environment reports. Retrieved from https://vcci.com.vn"
        ]
        references.extend(fallback_sources)
    
    # Add references to document
    for ref in references:
        ref_para = doc.add_paragraph(ref)
        ref_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        ref_para.paragraph_format.left_indent = Inches(0.3)
        ref_para.paragraph_format.hanging_indent = Inches(0.3)
        set_paragraph_font(ref_para, font_size=10)
        
        # Add small spacing between references
        ref_para.paragraph_format.space_after = Pt(3)
    
    # Add note about data sources
    doc.add_paragraph()
    note_para = doc.add_paragraph()
    note_para.add_run("Ghi chú về nguồn dữ liệu: ").bold = True
    note_para.add_run(f"Báo cáo này được tạo bằng AI ({api_provider} {model_used}) để phân tích và tổng hợp thông tin về thị trường {topic} tại {vietnamese_market}. "
                     f"AI được sử dụng để thu thập, phân tích và trình bày thông tin từ các nguồn công khai. "
                     f"Các nguồn tham khảo được trích xuất tự động từ quá trình phân tích và được sắp xếp theo tần suất sử dụng. "
                     f"Các thông tin và số liệu trong báo cáo phản ánh kiến thức và dữ liệu có sẵn của mô hình AI tại thời điểm tạo báo cáo ({datetime.now().strftime('%m/%Y')}). "
                     f"Người đọc nên xác minh thông tin với các nguồn chính thức và cập nhật để có dữ liệu mới nhất.")
    note_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_font(note_para, font_size=9)
    note_para.runs[0].italic = True
    note_para.runs[1].italic = True

def extract_sources_from_research(research_results):
    """Trích xuất các nguồn được đề cập trong nội dung research"""
    sources = set()
    
    # Common organizations and sources that might be mentioned in AI responses
    source_patterns = [
        # Vietnamese official sources
        r'Tổng cục Thống kê(?:\s+Việt\s+Nam)?',
        r'Bộ (?:Kế hoạch|Tài chính|Công Thương|Y tế|Giáo dục)',
        r'VCCI|Vietnam Chamber of Commerce',
        r'FPT|Viettel|VNPT',
        r'Ngân hàng Nhà nước',
        
        # International sources
        r'World Bank|Ngân hàng Thế giới',
        r'IMF|International Monetary Fund',
        r'ADB|Asian Development Bank',
        r'McKinsey|Deloitte|PwC|KPMG',
        r'Nielsen|Euromonitor',
        r'Statista',
        
        # Industry-specific
        r'VAMA|Vietnam Automobile',
        r'VINASA|Vietnam Software',
        r'VFA|Vietnam Food Association'
    ]
    
    # Search through all research content
    for layer in research_results:
        for category in layer.get('categories', []):
            for question in category.get('questions', []):
                # Check Layer 3 content
                content = question.get('layer3_content', '')
                if content:
                    for pattern in source_patterns:
                        matches = re.findall(pattern, content, re.IGNORECASE)
                        for match in matches:
                            sources.add(f"{match} - referenced in AI analysis")
                
                # Check Layer 4 comprehensive content
                if question.get('layer4_comprehensive_report'):
                    comp_content = question['layer4_comprehensive_report'].get('comprehensive_content', '')
                    if comp_content:
                        for pattern in source_patterns:
                            matches = re.findall(pattern, comp_content, re.IGNORECASE)
                            for match in matches:
                                sources.add(f"{match} - referenced in AI analysis")
    
    return list(sources)[:7]  # Return max 7 sources

def get_credible_industry_sources(topic, market):
    """Lấy nguồn uy tín theo ngành và thị trường"""
    topic_lower = topic.lower()
    sources = []
    current_year = datetime.now().year
    
    # Always include these for Vietnam market
    if 'việt nam' in market.lower() or 'vietnam' in market.lower():
        sources.extend([
            f"Vietnam Economic Times. ({current_year}). Industry Analysis Reports. Retrieved from https://vneconomictimes.com",
            f"Vietnam Investment Review. ({current_year}). Market Intelligence Reports. Retrieved from https://vir.com.vn"
        ])
    
    # Technology/Digital
    if any(keyword in topic_lower for keyword in ['technology', 'digital', 'ai', 'tech', 'công nghệ']):
        sources.extend([
            f"Vietnam ICT Statistics. ({current_year}). Ministry of Information and Communications. Retrieved from https://mic.gov.vn",
            f"VINASA Technology Reports. ({current_year}). Vietnam Software Association. Retrieved from https://vinasa.org.vn"
        ])
    
    # Automotive
    elif any(keyword in topic_lower for keyword in ['automotive', 'vehicle', 'ô tô', 'xe']):
        sources.extend([
            f"VAMA Industry Statistics. ({current_year}). Vietnam Automobile Manufacturers Association. Retrieved from https://vama.org.vn",
            f"Vietnam Ministry of Transport. ({current_year}). Transport Statistics. Retrieved from https://mt.gov.vn"
        ])
    
    # Default business sources
    else:
        sources.extend([
            f"Vietnam Business Portal. ({current_year}). Ministry of Planning and Investment. Retrieved from https://business.gov.vn",
            f"VCCI Business Reports. ({current_year}). Vietnam Chamber of Commerce. Retrieved from https://vcci.com.vn"
        ])
    
    return sources[:3]  # Return max 3 additional sources

def generate_ai_executive_summary(data):
    """Generate executive summary using AI based on all research questions and findings"""
    try:
        # Import here to avoid circular import
        from openai_market_research import OpenAIMarketResearch
        
        # Get API key from data or environment
        api_key = data.get('api_key') or os.getenv('OPENAI_API_KEY')
        if not api_key:
            return None
        
        # Initialize AI client
        researcher = OpenAIMarketResearch(
            api_key=api_key,
            industry=data.get('industry', 'Unknown'),
            market=data.get('market', 'Vietnam')
        )
        
        # Collect all main questions from research
        all_questions = []
        research_summary = ""
        
        for layer in data.get('research_results', []):
            layer_name = layer.get('layer_name', '')
            for category in layer.get('categories', []):
                category_name = category.get('category_name', '')
                for question in category.get('questions', []):
                    main_question = question.get('main_question', '')
                    if main_question:
                        all_questions.append(f"• {main_question}")
                        
                        # Get key insights from Layer 4 or Layer 3
                        layer4_report = question.get('layer4_comprehensive_report', {})
                        if layer4_report:
                            content = layer4_report.get('comprehensive_content', '')
                            # Extract first 2 sentences as key insight
                            sentences = content.split('.')[:2]
                            if sentences:
                                key_insight = '. '.join(sentences).strip()
                                if len(key_insight) > 50:
                                    research_summary += f"\n- {layer_name}/{category_name}: {key_insight[:200]}..."
                        elif question.get('layer3_content'):
                            content = question.get('layer3_content', '')
                            sentences = content.split('.')[:1]
                            if sentences:
                                key_insight = sentences[0].strip()
                                if len(key_insight) > 50:
                                    research_summary += f"\n- {layer_name}/{category_name}: {key_insight[:150]}..."
        
        questions_text = "\n".join(all_questions)
        
        # Create AI prompt for executive summary
        prompt = f"""Bạn là chuyên gia tư vấn chiến lược, viết tóm tắt điều hành (Executive Summary) cho báo cáo nghiên cứu thị trường.

THÔNG TIN NGHIÊN CỨU:
- Ngành: {data.get('industry', 'N/A')}
- Thị trường: {data.get('market', 'N/A')}
- Mục đích: {data.get('purpose', 'Phân tích thị trường và cơ hội kinh doanh')}

CÁC CÂU HỎI NGHIÊN CỨU ĐÃ ĐƯỢC PHÂN TÍCH:
{questions_text}

KEY INSIGHTS TỪ NGHIÊN CỨU:
{research_summary}

VIẾT TÓM TẮT ĐIỀU HÀNH theo 5 phần:

**1. 🎯 MỤC TIÊU/MỤC ĐÍCH (80-100 từ)**
Tóm tắt mục tiêu nghiên cứu và tại sao quan trọng

**2. 🌍 PHẠM VI/BỐI CẢNH (80-100 từ)**  
Phạm vi nghiên cứu, thị trường, phương pháp

**3. 🔍 PHÁT HIỆN CHÍNH (120-150 từ)**
3-4 insight quan trọng nhất từ nghiên cứu

**4. 🚀 ĐỀ XUẤT HÀNH ĐỘNG (120-150 từ)**
5-6 khuyến nghị cụ thể dựa trên findings

**5. 📈 TÁC ĐỘNG KỲ VỌNG (100-120 từ)**
Lợi ích và impact khi áp dụng các đề xuất

**YÊU CẦU:**
- Viết chuyên nghiệp, ngắn gọn, actionable
- Dựa trên insights thực từ nghiên cứu
- Tránh chung chung, focus vào specific findings
- Kết thúc mỗi phần bằng dấu chấm
- Format: Mỗi phần là một đoạn văn liền mạch

**CHỈ TRẢ VỀ NỘI DUNG 5 PHẦN, KHÔNG CÓ GIẢI THÍCH HAY INTRO**"""

        # Get AI-generated summary
        ai_summary = researcher.call_openai_api(prompt)
        
        return ai_summary
        
    except Exception as e:
        print(f"❌ Error generating AI executive summary: {e}")
        return None

def create_executive_summary(doc, data):
    """Tạo Executive Summary theo template 5 phần"""
    print("📋 Tạo Executive Summary...")
    
    # Get Vietnamese market name for consistency
    vietnamese_market = get_vietnamese_market_name(data.get('market', 'Việt Nam'))
    
    doc.add_page_break()
    
    # Executive Summary Header
    exec_heading = doc.add_heading('📋 TÓM TẮT ĐIỀU HÀNH (EXECUTIVE SUMMARY)', level=1)
    exec_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(exec_heading, font_size=18)
    
    # 1. Mục tiêu / Mục đích
    section1_heading = doc.add_heading('1. 🎯 MỤC TIÊU / MỤC ĐÍCH', level=2)
    set_paragraph_font(section1_heading, font_size=14)
    
    purpose_para = doc.add_paragraph()
    purpose_text = data.get('purpose', '')
    if purpose_text:
        # If custom purpose is provided, format it properly
        formatted_purpose = format_purpose_text(purpose_text)
        if not formatted_purpose.endswith('.'):
            formatted_purpose += '.'
        purpose_para.add_run(f"Báo cáo này nhằm {formatted_purpose}")
    else:
        purpose_para.add_run("Báo cáo này nhằm:")
        purpose_para.add_run(f"""
• Hiểu thị trường tổng thể: xu hướng, quy mô, tốc độ tăng trưởng của ngành {data.get('industry', 'N/A')} tại {vietnamese_market}
• Biết được mình đang định nhảy vào thị trường lớn hay nhỏ, chật chội hay đang mở
• Phân tích môi trường vĩ mô ảnh hưởng đến ngành (chính trị, kinh tế, công nghệ...)
• Dùng trước khi quyết định có nên vào thị trường này không, hoặc để thuyết phục nhà đầu tư""")
    
    purpose_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_font(purpose_para)
    
    # 2. Phạm vi / Bối cảnh
    section2_heading = doc.add_heading('2. 🌍 PHẠM VI / BỐI CẢNH', level=2)
    set_paragraph_font(section2_heading, font_size=14)
    
    scope_para = doc.add_paragraph()
    stats = calculate_statistics(data)
    scope_para.add_run(f"Tập trung vào lĩnh vực {data.get('industry', 'N/A')} tại thị trường {vietnamese_market} trong năm {datetime.now().year}, sử dụng phương pháp nghiên cứu phân tích đa tầng với {stats['total_questions']} câu hỏi nghiên cứu, áp dụng AI để thu thập và phân tích thông tin từ nhiều nguồn khác nhau.")
    scope_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_font(scope_para)
    
    # 3. Phát hiện chính
    section3_heading = doc.add_heading('3. 🔍 PHÁT HIỆN CHÍNH', level=2)
    set_paragraph_font(section3_heading, font_size=14)
    
    # Extract key insights from research content
    insights = extract_key_insights_for_summary(data)
    
    findings_para = doc.add_paragraph()
    findings_para.add_run("Chúng tôi nhận thấy rằng:")
    for i, insight in enumerate(insights[:4], 1):  # Top 4 insights for findings
        # Capitalize first letter of insight
        insight_text = insight['insight']
        if insight_text and len(insight_text) > 0:
            insight_text = insight_text[0].upper() + insight_text[1:] if len(insight_text) > 1 else insight_text.upper()
        findings_para.add_run(f"\n• {insight_text}")
    
    findings_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_font(findings_para)
    
    # 4. Đề xuất hành động
    section4_heading = doc.add_heading('4. 🚀 ĐỀ XUẤT HÀNH ĐỘNG', level=2)
    set_paragraph_font(section4_heading, font_size=14)
    
    action_para = doc.add_paragraph()
    action_para.add_run("Dựa trên kết quả nghiên cứu, chúng tôi đề xuất:")
    
    # Generate dynamic recommendations based on research content
    recommendations = generate_recommendations_from_research(data)
    
    if recommendations:
        for rec in recommendations[:5]:  # Top 5 recommendations
            action_para.add_run(f"\n• {rec}")
    else:
        # Fallback to generic recommendations
        action_para.add_run(f"""
• Phát triển chiến lược phù hợp với xu hướng thị trường đã xác định trong nghiên cứu
• Tập trung vào các cơ hội được nhận diện qua phân tích môi trường kinh doanh
• Đối phó với những thách thức chính được chỉ ra trong báo cáo
• Xây dựng năng lực cạnh tranh dựa trên insights từ phân tích ngành
• Thiết lập hệ thống giám sát để theo dõi sự thay đổi của các yếu tố được nghiên cứu""")
    
    action_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_font(action_para)
    
    # 5. Tác động kỳ vọng
    section5_heading = doc.add_heading('5. 📈 TÁC ĐỘNG KỲ VỌNG', level=2)
    set_paragraph_font(section5_heading, font_size=14)
    
    impact_para = doc.add_paragraph()
    
    # Generate dynamic impact based on research insights
    impact_insights = extract_impact_insights(data)
    
    if impact_insights:
        impact_para.add_run("Việc thực hiện các đề xuất trên dự kiến sẽ mang lại:")
        for impact in impact_insights[:3]:  # Top 3 impacts
            impact_para.add_run(f"\n• {impact}")
        
        impact_para.add_run(f"\n\nTổng thể, điều này sẽ giúp doanh nghiệp nâng cao khả năng cạnh tranh trong ngành {data.get('industry', 'N/A')} và thích ứng tốt hơn với môi trường kinh doanh năng động tại {vietnamese_market}.")
    else:
        # Fallback to more generic but still dynamic text
        impact_para.add_run(f"Việc áp dụng các insights từ nghiên cứu này sẽ giúp doanh nghiệp nâng cao vị thế cạnh tranh trong ngành {data.get('industry', 'N/A')}, tăng cường khả năng thích ứng với thay đổi thị trường tại {vietnamese_market}, và tối ưu hóa hiệu quả kinh doanh. Dự kiến sẽ cải thiện đáng kể khả năng ra quyết định chiến lược và tạo ra lợi thế cạnh tranh bền vững.")
    
    impact_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_font(impact_para)
    
    # Footer info
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.add_run('📅 Báo cáo được tạo tự động bởi Market Research Automation System').italic = True
    footer_para.add_run(f'\n⏰ Ngày tạo: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(footer_para, font_size=9)

def extract_key_insights_for_summary(data):
    """Trích xuất key insights để tạo executive summary phần phát hiện chính"""
    insights = []
    
    for layer in data.get('research_results', []):
        layer_name = layer.get('layer_name', '')
        
        for category in layer.get('categories', []):
            category_name = category.get('category_name', '')
            
            for question in category.get('questions', []):
                # Lấy từ comprehensive report trước
                layer4_comprehensive = question.get('layer4_comprehensive_report', {})
                if layer4_comprehensive:
                    content = layer4_comprehensive.get('comprehensive_content', '')
                    if content:
                        # Extract key trends and opportunities
                        sentences = content.split('.')
                        for sentence in sentences:
                            if any(keyword in sentence.lower() for keyword in ['xu hướng', 'cơ hội', 'thách thức', 'tăng trưởng', 'phát triển']):
                                insight = sentence.strip()
                                if len(insight) > 50:  # Only meaningful insights
                                    insights.append({
                                        'category': f"{layer_name} - {category_name}",
                                        'insight': insight[:150] + "..." if len(insight) > 150 else insight
                                    })
                                    break
                # Fallback to layer 3
                elif question.get('layer3_content'):
                    content = question.get('layer3_content', '')
                    sentences = content.split('.')
                    for sentence in sentences:
                        if any(keyword in sentence.lower() for keyword in ['quan trọng', 'chính', 'đáng chú ý', 'nổi bật']):
                            insight = sentence.strip()
                            if len(insight) > 50:
                                insights.append({
                                    'category': f"{layer_name} - {category_name}",
                                    'insight': insight[:120] + "..." if len(insight) > 120 else insight
                                })
                                break
    
    return insights[:6]  # Top 6 insights for summary

def generate_recommendations_from_research(data):
    """Generate dynamic recommendations based on research content"""
    recommendations = []
    
    for layer in data.get('research_results', []):
        layer_name = layer.get('layer_name', '')
        
        for category in layer.get('categories', []):
            category_name = category.get('category_name', '')
            
            for question in category.get('questions', []):
                # Lấy từ comprehensive report trước
                layer4_comprehensive = question.get('layer4_comprehensive_report', {})
                if layer4_comprehensive:
                    content = layer4_comprehensive.get('comprehensive_content', '')
                    if content:
                        # Extract actionable recommendations
                        sentences = content.split('.')
                        for sentence in sentences:
                            # Look for actionable insights and recommendations
                            if any(keyword in sentence.lower() for keyword in [
                                'nên', 'cần', 'khuyến nghị', 'đề xuất', 'tăng cường', 
                                'phát triển', 'đầu tư', 'tập trung', 'xây dựng', 'thúc đẩy'
                            ]):
                                rec = sentence.strip()
                                if len(rec) > 30 and rec not in recommendations:  # Avoid duplicates
                                    recommendations.append(rec[:200] + "..." if len(rec) > 200 else rec)
                                    if len(recommendations) >= 8:  # Enough recommendations
                                        break
                
                # Fallback to layer 3 content
                elif question.get('layer3_content'):
                    content = question.get('layer3_content', '')
                    sentences = content.split('.')
                    for sentence in sentences:
                        if any(keyword in sentence.lower() for keyword in [
                            'nên', 'cần', 'quan trọng', 'chính', 'ưu tiên'
                        ]):
                            rec = sentence.strip()
                            if len(rec) > 30 and rec not in recommendations:
                                recommendations.append(rec[:150] + "..." if len(rec) > 150 else rec)
                                if len(recommendations) >= 8:
                                    break
    
    return recommendations[:6]  # Top 6 actionable recommendations

def extract_impact_insights(data):
    """Extract expected impact insights from research content"""
    impacts = []
    
    for layer in data.get('research_results', []):
        for category in layer.get('categories', []):
            for question in category.get('questions', []):
                # Check comprehensive reports first
                layer4_comprehensive = question.get('layer4_comprehensive_report', {})
                if layer4_comprehensive:
                    content = layer4_comprehensive.get('comprehensive_content', '')
                    if content:
                        sentences = content.split('.')
                        for sentence in sentences:
                            # Look for impact-related statements
                            if any(keyword in sentence.lower() for keyword in [
                                'tác động', 'ảnh hưởng', 'hiệu quả', 'kết quả', 'lợi ích',
                                'cải thiện', 'tăng trưởng', 'giảm thiểu', 'tối ưu', 'nâng cao'
                            ]):
                                impact = sentence.strip()
                                if len(impact) > 40 and impact not in impacts:
                                    impacts.append(impact[:180] + "..." if len(impact) > 180 else impact)
                                    if len(impacts) >= 5:
                                        break
                
                # Fallback to layer 3
                elif question.get('layer3_content'):
                    content = question.get('layer3_content', '')
                    sentences = content.split('.')
                    for sentence in sentences:
                        if any(keyword in sentence.lower() for keyword in [
                            'lợi ích', 'hiệu quả', 'cải thiện', 'tăng', 'giảm'
                        ]):
                            impact = sentence.strip()
                            if len(impact) > 40 and impact not in impacts:
                                impacts.append(impact[:150] + "..." if len(impact) > 150 else impact)
                                if len(impacts) >= 5:
                                    break
    
    return impacts[:4]  # Top 4 impact insights

def clean_comprehensive_content(content):
    """Remove numbering, section headers, and intro sentences from comprehensive content"""
    if not content:
        return content
    
    # Remove intro sentences like "Để trả lời chuyên sâu câu hỏi về..."
    cleaned_content = re.sub(r'^Để trả lời[^,]*,\s*', '', content, flags=re.MULTILINE)
    cleaned_content = re.sub(r'^Để trả lời[^.]*\.\s*', '', cleaned_content, flags=re.MULTILINE)
    cleaned_content = re.sub(r'^Nhằm trả lời[^,]*,\s*', '', cleaned_content, flags=re.MULTILINE)
    cleaned_content = re.sub(r'^Nhằm phân tích[^,]*,\s*', '', cleaned_content, flags=re.MULTILINE)
    
    # Remove section headers like "📊 PHÂN TÍCH HIỆN TRẠNG:", "⚡ DRIVERS & IMPACTS:", etc.
    cleaned_content = re.sub(r'\n*[📊⚡📈⚠️🚀]\s*[A-ZÁÀẢÃẠĂẮẰẲẴẶÂẤẦẨẪẬĐÉÈẺẼẸÊẾỀỂỄỆÍÌỈĨỊÓÒỎÕỌÔỐỒỔỖỘƠỚỜỞỠỢÚÙỦŨỤƯỨỪỬỮỰÝỲỶỸỴ\s&]+:\s*', '\n\n', cleaned_content)
    
    # Remove numbered sections like "1. SECTION:", "2. ANALYSIS:", etc.
    cleaned_content = re.sub(r'\n*\d+\.\s+[A-ZÁÀẢÃẠĂẮẰẲẴẶÂẤẦẨẪẬĐÉÈẺẼẸÊẾỀỂỄỆÍÌỈĨỊÓÒỎÕỌÔỐỒỔỖỘƠỚỜỞỠỢÚÙỦŨỤƯỨỪỬỮỰÝỲỶỸỴ\s]+:\s*', '\n\n', cleaned_content)
    
    # Remove **bold headers** at start of lines
    cleaned_content = re.sub(r'\n\*\*[^*]+\*\*\s*\(\d+-\d+\s+từ\)\s*\n', '\n\n', cleaned_content)
    
    # Remove intro phrases at the beginning
    intro_patterns = [
        r'^Trong bối cảnh này,\s*',
        r'^Chúng ta cần xem xét[^.]*\.\s*',
        r'^Việc phân tích[^.]*\.\s*'
    ]
    
    for pattern in intro_patterns:
        cleaned_content = re.sub(pattern, '', cleaned_content, flags=re.MULTILINE)
    
    # Clean up extra newlines
    cleaned_content = re.sub(r'\n{3,}', '\n\n', cleaned_content)
    
    # Remove leading/trailing whitespace
    cleaned_content = cleaned_content.strip()
    
    return cleaned_content

def create_comprehensive_word_report(json_file: str, output_file: str = None, use_vietnamese_filename: bool = True) -> str:
    """
    Tạo báo cáo Word toàn diện từ kết quả nghiên cứu JSON
    
    Args:
        json_file: Đường dẫn file JSON
        output_file: Đường dẫn file output (optional)
        use_vietnamese_filename: Sử dụng tên file tiếng Việt thân thiện hay technical name
    """
    
    print("📄 Tạo trang bìa...")
    # Đọc dữ liệu
    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    if not output_file:
        # Extract base info for filename
        base_name = os.path.basename(json_file).replace('layer3_research_', '').replace('.json', '')
        
        if use_vietnamese_filename:
            # Tạo tên file tiếng Việt thân thiện
            metadata = data.get('research_metadata', {})
            industry = metadata.get('industry', 'Unknown')
            market = metadata.get('market', 'Vietnam')
            timestamp = metadata.get('research_timestamp', '').replace(':', '').replace(' ', '_').replace('-', '')
            
            # Clean industry name for filename
            industry_clean = re.sub(r'[^\w\s-]', '', industry)
            industry_clean = re.sub(r'\s+', '_', industry_clean)
            
            # Create friendly Vietnamese filename
            output_file = f"Báo_cáo_nghiên_cứu_thị_trường_{industry_clean}_{timestamp[:8]}.docx"
        else:
            # Use technical naming convention
            output_file = f"{base_name}_comprehensive_report.docx"
        
        # Ensure output directory
        output_dir = os.path.dirname(json_file) if '/' in json_file else 'output'
        output_file = os.path.join(output_dir, output_file)
    
    # Create document
    doc = Document()
    add_custom_styles(doc)
    
    # ===== TRANG BÌA =====
    print("📄 Tạo trang bìa...")
    
    title = doc.add_heading('BÁO CÁO NGHIÊN CỨU THỊ TRƯỜNG', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(title, font_size=18)
    
    doc.add_paragraph()
    
    # Thông tin cơ bản với Table
    info_heading = doc.add_heading('📊 THÔNG TIN TỔNG QUAN', level=2)
    info_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(info_heading, font_size=14)
    
    create_info_table(doc, data)
    
    doc.add_paragraph()
    
    # Purpose
    if data.get('purpose'):
        purpose_heading = doc.add_heading('🎯 MỤC ĐÍCH NGHIÊN CỨU', level=2)
        purpose_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_font(purpose_heading, font_size=14)
        
        purpose_para = doc.add_paragraph(data['purpose'])
        purpose_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_font(purpose_para)
    
    doc.add_page_break()
    
    # ===== NỘI DUNG CHÍNH =====
    print("🔍 Tạo nội dung báo cáo chi tiết...")
    
    results_heading = doc.add_heading('📋 KẾT QUẢ NGHIÊN CỨU CHI TIẾT', level=1)
    results_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(results_heading, font_size=16)
    
    for layer_idx, layer in enumerate(data.get('research_results', []), 1):
        layer_name = layer.get('layer_name', '')
        
        print(f"  🔥 Xử lý Layer: {layer_name}")
        
        # Layer heading
        layer_heading = doc.add_heading(f"{layer_idx}. {layer_name.upper()}", level=1)
        set_paragraph_font(layer_heading, font_size=16)
        
        for cat_idx, category in enumerate(layer.get('categories', []), 1):
            category_name = category.get('category_name', '')
            
            print(f"    📋 Category: {category_name}")
            
            # Category heading
            category_heading = doc.add_heading(f"{cat_idx}. {category_name}", level=2)
            set_paragraph_font(category_heading, font_size=14)
            
            # Check if this is Layer 3 comprehensive category analysis
            layer3_comprehensive_category = category.get('layer3_comprehensive_category', {})
            if layer3_comprehensive_category:
                print(f"    🎯 Layer 3 Comprehensive Category Analysis: {category_name}")
                
                # Display Layer 3 comprehensive analysis for entire category
                comprehensive_content = layer3_comprehensive_category.get('comprehensive_content', '')
                if comprehensive_content:
                    # Clean up content
                    cleaned_content = clean_comprehensive_content(comprehensive_content).strip()
                    
                    # Display analysis content directly without headers
                    comp_para = doc.add_paragraph(cleaned_content)
                    comp_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    set_paragraph_font(comp_para)
                
                # Timestamp only
                timestamp = layer3_comprehensive_category.get('analysis_timestamp', '')
                if timestamp:
                    time_para = doc.add_paragraph()
                    time_para.add_run(f"⏰ Phân tích: {timestamp}").italic = True
                    time_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    set_paragraph_font(time_para, font_size=9)
                
                # Skip individual question processing for Layer 3 mode
                continue
            
            # Layer 4 mode: Process individual questions (existing logic)
            for q_idx, question in enumerate(category.get('questions', []), 1):
                main_question = question.get('main_question', '')
                
                print(f"      ❓ Question: {main_question[:50]}...")
                
                # Main question heading
                question_heading = doc.add_heading(f"{q_idx}. {main_question}", level=3)
                set_paragraph_font(question_heading, font_size=12)
                
                # Layer 3 content (more concise)
                layer3_content = question.get('layer3_content', '')
                if layer3_content:
                    # Clean up and make more concise
                    cleaned_content = layer3_content.strip()
                    
                    layer3_para = doc.add_paragraph()
                    layer3_para.add_run('📋 PHÂN TÍCH TỔNG QUAN:').bold = True
                    layer3_para.add_run('\n' + cleaned_content)
                    layer3_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    set_paragraph_font(layer3_para)
                
                # Layer 4 comprehensive report (ưu tiên)
                layer4_comprehensive = question.get('layer4_comprehensive_report', {})
                if layer4_comprehensive:
                    print(f"        🔍 Có comprehensive report")
                    
                    comprehensive_content = layer4_comprehensive.get('comprehensive_content', '')
                    if comprehensive_content:
                        # Clean up content and make more structured
                        cleaned_comp_content = clean_comprehensive_content(comprehensive_content).strip()
                        
                        comp_para = doc.add_paragraph()
                        comp_para.add_run('🎯 PHÂN TÍCH CHUYÊN SÂU:').bold = True
                        comp_para.add_run('\n' + cleaned_comp_content)
                        comp_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        set_paragraph_font(comp_para)
                    
                    # Timestamp
                    timestamp = layer4_comprehensive.get('enhancement_timestamp', '')
                    if timestamp:
                        time_para = doc.add_paragraph()
                        time_para.add_run(f"⏰ Cập nhật: {timestamp}").italic = True
                        time_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        set_paragraph_font(time_para, font_size=9)
                
                # Layer 4 individual enhancements (fallback)
                elif question.get('layer4_enhancements', {}):
                    layer4_enhancements = question.get('layer4_enhancements', {})
                    print(f"        🎯 Có {len(layer4_enhancements)} individual enhancements")
                    
                    layer4_heading = doc.add_paragraph()
                    layer4_heading.add_run('🎯 PHÂN TÍCH CHI TIẾT:').bold = True
                    
                    for sub_question, enhancement_data in layer4_enhancements.items():
                        # Sub-question
                        sub_q_para = doc.add_paragraph()
                        sub_q_para.add_run(f"• {sub_question}").bold = True
                        set_paragraph_font(sub_q_para)
                        
                        # Enhanced content
                        enhanced_content = enhancement_data.get('enhanced_content', '')
                        if enhanced_content:
                            enh_para = doc.add_paragraph(enhanced_content)
                            enh_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            enh_para.paragraph_format.left_indent = Inches(0.3)
                            set_paragraph_font(enh_para)
                        
                        # Timestamp
                        timestamp = enhancement_data.get('enhancement_timestamp', '')
                        if timestamp:
                            time_para = doc.add_paragraph()
                            time_para.add_run(f"⏰ Cập nhật: {timestamp}").italic = True
                            time_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            set_paragraph_font(time_para, font_size=9)
    
    # ===== REFERENCES SECTION =====
    create_references_section(doc, data)
    
    # ===== EXECUTIVE SUMMARY =====
    print("📋 Generating AI-powered Executive Summary...")
    
    # Try AI-generated summary first
    ai_summary = generate_ai_executive_summary(data)
    
    if ai_summary:
        # Add page break before executive summary
        doc.add_page_break()
        
        # Executive Summary Header
        exec_heading = doc.add_heading('📋 TÓM TẮT ĐIỀU HÀNH (EXECUTIVE SUMMARY)', level=1)
        exec_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_font(exec_heading, font_size=18)
        
        # Add AI-generated content
        ai_summary_para = doc.add_paragraph(ai_summary)
        ai_summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_font(ai_summary_para)
        
        # Footer info
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.add_run('📅 Báo cáo được tạo tự động bởi Market Research Automation System').italic = True
        footer_para.add_run(f'\n⏰ Ngày tạo: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
        footer_para.add_run(f'\n🤖 Executive Summary generated by AI based on research findings')
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_font(footer_para, font_size=9)
    else:
        # Fallback to original executive summary
        print("⚠️ AI summary failed, using fallback...")
        create_executive_summary(doc, data)
    
    # Save document
    print(f"💾 Lưu file: {output_file}")
    doc.save(output_file)
    print(f"✅ Đã tạo Word document: {output_file}")
    
    return output_file

def calculate_statistics(data):
    """Tính toán thống kê cho báo cáo"""
    stats = {
        'total_layers': 0,
        'total_categories': 0,
        'total_questions': 0,
        'comprehensive_reports': 0,
        'individual_enhancements': 0
    }
    
    research_results = data.get('research_results', [])
    stats['total_layers'] = len(research_results)
    
    for layer in research_results:
        categories = layer.get('categories', [])
        stats['total_categories'] += len(categories)
        
        for category in categories:
            questions = category.get('questions', [])
            stats['total_questions'] += len(questions)
            
            for question in questions:
                # Count comprehensive reports
                if question.get('layer4_comprehensive_report'):
                    stats['comprehensive_reports'] += 1
                
                # Count individual enhancements
                individual_enhancements = question.get('layer4_enhancements', {})
                stats['individual_enhancements'] += len(individual_enhancements)
    
    return stats

def find_latest_research_file():
    """Tìm file research mới nhất trong thư mục output"""
    output_dir = 'output'
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    files = []
    for file in os.listdir(output_dir):
        if file.startswith('layer3_research_') and file.endswith('.json'):
            files.append(os.path.join(output_dir, file))
    
    if not files:
        # Fallback: tìm trong thư mục gốc
        files = [f for f in os.listdir('.') if f.startswith('layer3_research_') and f.endswith('.json')]
        
    if not files:
        return None
    
    # Sắp xếp theo thời gian modification
    files.sort(key=lambda x: os.path.getmtime(x), reverse=True)
    return files[0]

def main():
    """Main function"""
    print("🔬 COMPREHENSIVE WORD EXPORT TOOL")
    print("="*50)
    
    # Tạo thư mục output nếu chưa có
    output_dir = 'output'
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"📁 Đã tạo thư mục: {output_dir}")
    
    # Tìm file research
    json_file = find_latest_research_file()
    
    if not json_file:
        print("❌ Không tìm thấy file research results!")
        print("💡 Hãy chạy Layer 3 research trước.")
        return
    
    print(f"📁 Sử dụng file: {json_file}")
    
    # Tạo output filename trong thư mục output
    base_name = os.path.basename(json_file).replace('layer3_research_', '').replace('.json', '')
    output_file = os.path.join(output_dir, f"Báo_cáo_nghiên_cứu_thị_trường_{base_name}.docx")
    
    # Export
    result = create_comprehensive_word_report(json_file, output_file)
    
    if result:
        print(f"\n🎉 HOÀN THÀNH!")
        print(f"📄 File Word: {result}")
        print(f"💡 Có thể mở file bằng: open \"{result}\"")
    else:
        print("❌ Export thất bại!")

def format_purpose_text(purpose_text):
    """Format purpose text to ensure proper capitalization for both bullet and dash formats"""
    if not purpose_text:
        return purpose_text
    
    # Split by common delimiters and capitalize each item
    lines = purpose_text.split('\n')
    formatted_lines = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Handle both bullet (•) and dash (-) formats
        if line.startswith('•') or line.startswith('-'):
            # Extract the text after the bullet/dash
            prefix = line[0]  # • or -
            text = line[1:].strip()
            if text:
                # Capitalize first letter
                text = text[0].upper() + text[1:] if len(text) > 1 else text.upper()
                formatted_lines.append(f"{prefix} {text}")
        else:
            # Regular text - capitalize first letter
            if line:
                line = line[0].upper() + line[1:] if len(line) > 1 else line.upper()
                formatted_lines.append(line)
    
    return '\n'.join(formatted_lines)

if __name__ == "__main__":
    main() 