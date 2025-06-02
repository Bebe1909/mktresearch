#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word Exporter cho Market Research Results
Chuyển đổi JSON results thành Word document có cấu trúc theo layers
"""

import json
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

class MarketResearchWordExporter:
    def __init__(self):
        """Khởi tạo Word exporter"""
        self.doc = Document()
        self.setup_styles()
        
    def setup_styles(self):
        """Tạo custom styles cho Word document"""
        
        # Title style
        title_style = self.doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Arial'
        title_font.size = Pt(24)
        title_font.bold = True
        title_font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(20)
        
        # Subtitle style
        subtitle_style = self.doc.styles.add_style('CustomSubtitle', WD_STYLE_TYPE.PARAGRAPH)
        subtitle_font = subtitle_style.font
        subtitle_font.name = 'Arial'
        subtitle_font.size = Pt(16)
        subtitle_font.bold = True
        subtitle_font.color.rgb = RGBColor(0, 102, 204)  # Blue
        subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_style.paragraph_format.space_after = Pt(15)
        
        # Layer heading style
        layer_style = self.doc.styles.add_style('LayerHeading', WD_STYLE_TYPE.PARAGRAPH)
        layer_font = layer_style.font
        layer_font.name = 'Arial'
        layer_font.size = Pt(18)
        layer_font.bold = True
        layer_font.color.rgb = RGBColor(46, 134, 171)  # Teal blue
        layer_style.paragraph_format.space_before = Pt(20)
        layer_style.paragraph_format.space_after = Pt(10)
        
        # Category heading style
        category_style = self.doc.styles.add_style('CategoryHeading', WD_STYLE_TYPE.PARAGRAPH)
        category_font = category_style.font
        category_font.name = 'Arial'
        category_font.size = Pt(14)
        category_font.bold = True
        category_font.color.rgb = RGBColor(162, 59, 114)  # Purple
        category_style.paragraph_format.left_indent = Inches(0.3)
        category_style.paragraph_format.space_before = Pt(15)
        category_style.paragraph_format.space_after = Pt(8)
        
        # Question heading style
        question_style = self.doc.styles.add_style('QuestionHeading', WD_STYLE_TYPE.PARAGRAPH)
        question_font = question_style.font
        question_font.name = 'Arial'
        question_font.size = Pt(12)
        question_font.bold = True
        question_font.color.rgb = RGBColor(241, 143, 1)  # Orange
        question_style.paragraph_format.left_indent = Inches(0.6)
        question_style.paragraph_format.space_before = Pt(12)
        question_style.paragraph_format.space_after = Pt(6)
        
        # Sub question style
        sub_question_style = self.doc.styles.add_style('SubQuestion', WD_STYLE_TYPE.PARAGRAPH)
        sub_question_font = sub_question_style.font
        sub_question_font.name = 'Arial'
        sub_question_font.size = Pt(11)
        sub_question_font.bold = True
        sub_question_font.color.rgb = RGBColor(199, 62, 29)  # Red-orange
        sub_question_style.paragraph_format.left_indent = Inches(0.9)
        sub_question_style.paragraph_format.space_before = Pt(8)
        sub_question_style.paragraph_format.space_after = Pt(4)
        
        # Answer text style
        answer_style = self.doc.styles.add_style('AnswerText', WD_STYLE_TYPE.PARAGRAPH)
        answer_font = answer_style.font
        answer_font.name = 'Arial'
        answer_font.size = Pt(10)
        answer_style.paragraph_format.left_indent = Inches(1.2)
        answer_style.paragraph_format.space_after = Pt(10)
        answer_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Summary box style
        summary_style = self.doc.styles.add_style('SummaryBox', WD_STYLE_TYPE.PARAGRAPH)
        summary_font = summary_style.font
        summary_font.name = 'Arial'
        summary_font.size = Pt(10)
        summary_style.paragraph_format.left_indent = Inches(0.3)
        summary_style.paragraph_format.right_indent = Inches(0.3)
        summary_style.paragraph_format.space_before = Pt(10)
        summary_style.paragraph_format.space_after = Pt(10)

    def add_cover_page(self, data):
        """Tạo trang bìa"""
        
        # Title
        title = self.doc.add_paragraph("BÁO CÁO NGHIÊN CỨU THỊ TRƯỜNG", style='CustomTitle')
        
        # Add some space
        self.doc.add_paragraph()
        
        # Industry & Market
        industry = data.get('industry', 'N/A')
        market = data.get('market', 'N/A')
        subtitle = f"{industry.upper()} - THỊ TRƯỜNG {market.upper()}"
        self.doc.add_paragraph(subtitle, style='CustomSubtitle')
        
        # Add space
        self.doc.add_paragraph()
        
        # Purpose section
        purpose_heading = self.doc.add_paragraph()
        purpose_heading.add_run("MỤC ĐÍCH NGHIÊN CỨU:").bold = True
        purpose_heading.style.font.size = Pt(14)
        
        purpose = data.get('purpose', 'Không có mô tả')
        purpose_para = self.doc.add_paragraph(purpose)
        purpose_para.style.font.name = 'Arial'
        purpose_para.style.font.size = Pt(11)
        
        # Add space
        self.doc.add_paragraph()
        
        # Statistics section
        stats = self.calculate_statistics(data)
        stats_heading = self.doc.add_paragraph()
        stats_heading.add_run("THỐNG KÊ TỔNG QUAN:").bold = True
        stats_heading.style.font.size = Pt(14)
        
        # Create statistics table
        table = self.doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        stats_data = [
            ('Số lượng Layers:', str(stats['total_layers'])),
            ('Số lượng Categories:', str(stats['total_categories'])),
            ('Số lượng Questions:', str(stats['total_questions'])),
            ('Số lượng Sub-questions:', str(stats['total_sub_questions'])),
            ('Ngày tạo báo cáo:', datetime.now().strftime('%d/%m/%Y %H:%M'))
        ]
        
        for i, (label, value) in enumerate(stats_data):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            
            # Format cells
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style.font.name = 'Arial'
                    paragraph.style.font.size = Pt(10)
            
            # Make first column bold
            row.cells[0].paragraphs[0].runs[0].bold = True
        
        # Page break
        self.doc.add_page_break()

    def calculate_statistics(self, data):
        """Tính toán thống kê"""
        stats = {
            'total_layers': 0,
            'total_categories': 0,
            'total_questions': 0,
            'total_sub_questions': 0
        }
        
        research_results = data.get('research_results', [])
        stats['total_layers'] = len(research_results)
        
        for layer in research_results:
            categories = layer.get('categories', [])
            stats['total_categories'] += len(categories)
            
            for category in categories:
                questions = category.get('questions', [])
                stats['total_questions'] += len(questions)
                
                for question in questions:
                    sub_answers = question.get('sub_answers', [])
                    stats['total_sub_questions'] += len(sub_answers)
        
        return stats

    def add_table_of_contents(self, data):
        """Tạo mục lục"""
        
        toc_title = self.doc.add_paragraph("MỤC LỤC", style='CustomTitle')
        
        # Add TOC entries
        research_results = data.get('research_results', [])
        for i, layer in enumerate(research_results, 1):
            layer_name = layer.get('layer_name', '')
            
            # Layer entry
            layer_entry = self.doc.add_paragraph()
            layer_run = layer_entry.add_run(f"{i}. {layer_name}")
            layer_run.bold = True
            layer_run.font.size = Pt(12)
            
            # Category entries
            categories = layer.get('categories', [])
            for j, category in enumerate(categories, 1):
                category_name = category.get('category_name', '')
                cat_entry = self.doc.add_paragraph()
                cat_entry.paragraph_format.left_indent = Inches(0.3)
                cat_run = cat_entry.add_run(f"{i}.{j} {category_name}")
                cat_run.font.size = Pt(11)
        
        # Page break
        self.doc.add_page_break()

    def add_layer_content(self, layer_data):
        """Thêm nội dung cho một layer"""
        
        layer_name = layer_data.get('layer_name', '')
        
        # Layer title
        self.doc.add_paragraph(f"LAYER: {layer_name}", style='LayerHeading')
        
        # Layer summary
        categories = layer_data.get('categories', [])
        total_questions = sum(len(cat.get('questions', [])) for cat in categories)
        total_sub_questions = sum(
            len(q.get('sub_answers', []))
            for cat in categories
            for q in cat.get('questions', [])
        )
        
        summary_para = self.doc.add_paragraph(style='SummaryBox')
        summary_para.add_run(f"Tổng quan Layer {layer_name}:").bold = True
        summary_para.add_run(f"\n• Số categories: {len(categories)}")
        summary_para.add_run(f"\n• Số main questions: {total_questions}")
        summary_para.add_run(f"\n• Số sub-questions: {total_sub_questions}")
        
        # Categories content
        for category in categories:
            self.add_category_content(category)
        
        # Page break after each layer
        self.doc.add_page_break()

    def add_category_content(self, category_data):
        """Thêm nội dung cho một category"""
        
        category_name = category_data.get('category_name', '')
        
        # Category title
        self.doc.add_paragraph(f"Category: {category_name}", style='CategoryHeading')
        
        # Questions content
        questions = category_data.get('questions', [])
        for question in questions:
            self.add_question_content(question)

    def add_question_content(self, question_data):
        """Thêm nội dung cho một question"""
        
        main_question = question_data.get('main_question', '')
        
        # Question title
        question_para = self.doc.add_paragraph(f"❓ {main_question}", style='QuestionHeading')
        
        # Kiểm tra research standard
        research_standard = question_data.get('research_standard', 'Layer 4')
        
        if research_standard == 'Layer 3':
            # Hiển thị Layer 3 content
            layer3_content = question_data.get('layer3_content', '')
            if layer3_content:
                # Layer 3 content
                content_para = self.doc.add_paragraph(style='AnswerContent')
                content_para.add_run("📊 Layer 3 Analysis:").bold = True
                content_para.add_run(f"\n{layer3_content}")
                
                # Layer 4 enhancements (nếu có)
                layer4_enhancements = question_data.get('layer4_enhancements', {})
                if layer4_enhancements:
                    enhancement_para = self.doc.add_paragraph()
                    enhancement_run = enhancement_para.add_run("\n🔍 Layer 4 Detailed Enhancements:")
                    enhancement_run.bold = True
                    enhancement_run.font.color.rgb = RGBColor(204, 102, 0)  # Orange
                    
                    for sub_question, enhancement_data in layer4_enhancements.items():
                        enhanced_content = enhancement_data.get('enhanced_content', '')
                        timestamp = enhancement_data.get('enhancement_timestamp', '')
                        
                        # Sub-question title
                        sub_para = self.doc.add_paragraph()
                        sub_para.add_run(f"\n• {sub_question}").bold = True
                        if timestamp:
                            sub_para.add_run(f" (Enhanced: {timestamp})").italic = True
                        
                        # Enhanced content
                        enhancement_content_para = self.doc.add_paragraph(style='AnswerContent')
                        enhancement_content_para.add_run(enhanced_content)
                        
                        # Add separator
                        self.doc.add_paragraph("─" * 50)
            
        else:
            # Legacy Layer 4 structure (sub_answers)
            sub_answers = question_data.get('sub_answers', [])
            
            if sub_answers:
                for sub_answer in sub_answers:
                    sub_question = sub_answer.get('sub_question', '')
                    answer = sub_answer.get('answer', '')
                    
                    # Sub-question
                    sub_para = self.doc.add_paragraph()
                    sub_run = sub_para.add_run(f"• {sub_question}")
                    sub_run.bold = True
                    sub_run.font.color.rgb = RGBColor(102, 102, 102)  # Gray
                    
                    # Answer
                    if answer and not answer.startswith('Lỗi:'):
                        answer_para = self.doc.add_paragraph(style='AnswerContent')
                        
                        # Truncate very long answers
                        if len(answer) > 3000:
                            answer = answer[:3000] + "...\n[Nội dung đã được rút gọn]"
                        
                        answer_para.add_run(answer)
                    else:
                        error_para = self.doc.add_paragraph()
                        error_run = error_para.add_run("⚠️ Không có dữ liệu hoặc có lỗi trong quá trình xử lý")
                        error_run.italic = True
                        error_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                    
                    # Add some spacing
                    self.doc.add_paragraph()
            else:
                no_data_para = self.doc.add_paragraph()
                no_data_run = no_data_para.add_run("⚠️ Không có dữ liệu cho câu hỏi này")
                no_data_run.italic = True
                no_data_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange

    def clean_answer_text(self, text):
        """Làm sạch text cho Word"""
        if not text:
            return "Không có dữ liệu"
        
        text = str(text)
        
        # Remove markdown formatting
        text = text.replace('**', '')
        text = text.replace('##', '')
        text = text.replace('#', '')
        
        # Clean up excessive newlines
        text = text.replace('\n\n\n', '\n\n')
        
        # Limit length if too long
        if len(text) > 3000:
            text = text[:3000] + "... (nội dung đã được rút gọn)"
        
        return text

    def export_to_word(self, json_file, output_file=None):
        """Export JSON results to Word document"""
        
        # Read JSON data
        with open(json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Generate output filename if not provided
        if output_file is None:
            industry = data.get('industry', 'unknown').replace(' ', '_')
            market = data.get('market', 'unknown')
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_file = f"market_research_{industry}_{market}_{timestamp}.docx"
        
        # Build document
        # Cover page
        self.add_cover_page(data)
        
        # Table of contents
        self.add_table_of_contents(data)
        
        # Layer content
        research_results = data.get('research_results', [])
        for layer in research_results:
            self.add_layer_content(layer)
        
        # Save document
        self.doc.save(output_file)
        
        print(f"✅ Đã xuất Word document thành công: {output_file}")
        return output_file

def main():
    """Test function"""
    exporter = MarketResearchWordExporter()
    
    # Test với demo data
    if Path("demo_research_result.json").exists():
        output_file = exporter.export_to_word("demo_research_result.json")
        print(f"📄 Word document demo đã được tạo: {output_file}")
    else:
        print("❌ Không tìm thấy demo_research_result.json")
        print("💡 Chạy 'python demo_flow_without_api.py' trước để tạo demo data")

if __name__ == "__main__":
    main() 